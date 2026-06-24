"""Sección de flujo de carga nodal en la memoria SEC."""
from pathlib import Path


def _tmp():
    d = Path(__file__).resolve().parent / ".tmp_flujo_nodal"
    d.mkdir(parents=True, exist_ok=True)
    return str(d)

def _datos_run():
    return {
        "project_id": "TEST-FN", "revision": "R1",
        "timestamp": "2026-06-23T12:00:00+00:00", "perfil": "industrial",
        "norma": "MM2", "n_ok": 1, "n_fallas": 0, "max_dv_pct": 1.0,
        "max_icc_ka": 30.0, "status": "OK",
        "trafo_z_ohm": 0.005, "tension_sistema_v": 380.0,
        "cadena": [
            {"nombre": "G0", "upstream": "", "nivel": 0, "In_A": 1600, "curva": "ETU600", "Icc_kA": 30.0},
            {"nombre": "G1", "upstream": "G0", "nivel": 1, "In_A": 630, "curva": "ETU320", "Icc_kA": 10.0},
            {"nombre": "C2", "upstream": "G1", "nivel": 2, "In_A": 160, "curva": "C", "Icc_kA": 5.0},
        ],
    }

def _circuitos():
    return [
        {"nombre": "L1", "S_mm2": 25.0, "sistema": "3F", "I_diseno": 100.0, "cos_phi": 0.9,
         "L_m": 20.0, "paralelos": 1, "dv_pct": 0.3, "icc_ka": 6.0, "estado": "OK", "norma": "MM2"},
    ]

def test_memoria_incluye_seccion_flujo_nodal():
    from docx import Document
    from reporteria_sec import generar_memoria_docx
    ruta = generar_memoria_docx(_datos_run(), _circuitos(), _tmp())
    doc = Document(ruta)
    texto = "\n".join(p.text for p in doc.paragraphs)
    assert "Flujo de Carga Nodal" in texto

def test_memoria_flujo_nodal_tabla_por_barra():
    from docx import Document
    from reporteria_sec import generar_memoria_docx
    ruta = generar_memoria_docx(_datos_run(), _circuitos(), _tmp())
    doc = Document(ruta)
    encontrada = any(
        any("V (pu)" in cell.text for cell in tabla.rows[0].cells)
        for tabla in doc.tables
    )
    assert encontrada

def test_memoria_flujo_nodal_sin_cadena_omite():
    from docx import Document
    from reporteria_sec import generar_memoria_docx
    datos = _datos_run()
    datos["cadena"] = []
    ruta = generar_memoria_docx(datos, _circuitos(), _tmp())
    doc = Document(ruta)
    texto = "\n".join(p.text for p in doc.paragraphs)
    assert "sin cadena" in texto.lower()

def test_json_epc_incluye_flujo_nodal():
    import json
    from reporteria_sec import exportar_json_epc
    ruta = exportar_json_epc(_datos_run(), _tmp(), circuitos=_circuitos())
    with open(ruta, encoding="utf-8") as fh:
        payload = json.load(fh)
    assert "flujo_nodal" in payload
    fn = payload["flujo_nodal"]
    assert "buses" in fn and "perdidas_totales_kW" in fn
    assert any(b["id"] == "TRAFO" for b in fn["buses"])
