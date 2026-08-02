import os
import json
import re
import sqlite3
import uuid
from pathlib import Path

from persistencia import registrar_ejecucion
from reporteria_sec import (
    exportar_json_epc,
    generar_desde_run_id,
    generar_memoria_docx,
    generar_memoria_sec,
    generar_reporte_pdf,
    verificar_completitud_parametros,
)


def _tmp_dir() -> Path:
    base = Path(__file__).resolve().parent / ".tmp_reporteria_sec_tests"
    base.mkdir(parents=True, exist_ok=True)
    return base


def _ruta_db_local(nombre: str) -> str:
    return str(_tmp_dir() / f"{nombre}_{uuid.uuid4().hex}.db")


def _datos_run_base():
    return {
        "project_id": "LEO-ARICA",
        "revision": "SEC1",
        "timestamp": "2026-04-19T12:00:00+00:00",
        "perfil": "industrial",
        "norma": "AWG",
        "n_ok": 2,
        "n_fallas": 0,
        "max_dv_pct": 1.5,
        "max_icc_ka": 20.0,
        "status": "OK",
    }


def _circuitos_base():
    return [
        {
            "nombre": "C-01",
            "conductor": "6AWG",
            "S_mm2": 13.3,
            "I_diseno": 42.0,
            "I_max": 65.0,
            "cos_phi": 0.9,
            "L_m": 15.0,
            "paralelos": 1,
            "sistema": "3F",
            "dv_v": 1.0,
            "dv_pct": 0.26,
            "icc_ka": 10.5,
            "estado": "OK",
            "norma": "AWG",
            "observaciones": "sin observaciones",
        }
    ]


def test_generar_memoria_docx():
    ruta = generar_memoria_docx(_datos_run_base(), _circuitos_base(), str(_tmp_dir()))
    assert ruta.endswith(".docx")
    assert Path(ruta).exists()


def test_generar_reporte_pdf():
    ruta = generar_reporte_pdf(_datos_run_base(), _circuitos_base(), str(_tmp_dir()))
    assert ruta.endswith(".pdf")
    assert Path(ruta).exists()


def test_nombres_archivo():
    ruta = generar_memoria_docx(_datos_run_base(), _circuitos_base(), str(_tmp_dir()))
    nombre = Path(ruta).name
    assert re.match(r"^MEMORIA_LEO-ARICA_SEC1_\d{8}_\d{4}\.docx$", nombre)


def test_generar_desde_run_id():
    ruta_db = _ruta_db_local("desde_run")
    run_id = registrar_ejecucion(
        {
            **_datos_run_base(),
            "ruta_reporte_txt": "reporte.txt",
            "ruta_reporte_xlsx": "reporte.xlsx",
            "circuitos": _circuitos_base(),
        },
        ruta_db=ruta_db,
    )

    cwd = os.getcwd()
    os.chdir(str(_tmp_dir()))
    try:
        ruta_docx, ruta_pdf = generar_desde_run_id(run_id, ruta_db=ruta_db)
    finally:
        os.chdir(cwd)

    assert ruta_docx.endswith(".docx")
    assert ruta_pdf.endswith(".pdf")
    assert Path(ruta_docx).exists()
    assert Path(ruta_pdf).exists()

    with sqlite3.connect(ruta_db) as conn:
        rows = conn.execute(
            "SELECT report_type, file_path FROM run_reports WHERE run_id = ?",
            (run_id,),
        ).fetchall()
    tipos = {r[0] for r in rows}
    assert "DOCX" in tipos
    assert "PDF" in tipos


def test_fallo_silencioso():
    ruta_db = _ruta_db_local("silencioso")
    try:
        rutas = generar_desde_run_id("run-id-no-existe", ruta_db=ruta_db)
    except Exception as e:
        raise AssertionError(f"No debia lanzar excepcion: {e}")
    assert rutas == ("", "")


def test_gate_emision_bloquea_con_defaults():
    gate = verificar_completitud_parametros(
        {
            "ats": {"t_arranque_ge_ms": 10000.0},
            "motor": {"factor_arranque": 6.0},
        }
    )
    assert gate["apto_emision"] is False
    assert gate["nivel"] == "INCOMPLETO"
    assert len(gate["parametros_default"]) >= 1


def test_gate_emision_aprueba_sin_defaults():
    gate = verificar_completitud_parametros(
        {
            "ats": {"t_arranque_ge_ms": 9000.0},
            "motor": {"factor_arranque": 5.5},
        }
    )
    assert gate["apto_emision"] is True
    assert gate["nivel"] == "FINAL"


def test_memoria_borrador_incluye_lista_defaults():
    datos = {
        **_datos_run_base(),
        "ats": {"t_arranque_ge_ms": 10000.0},
    }
    ruta = generar_memoria_sec(datos, _circuitos_base(), str(_tmp_dir()), modo_emision="auto")
    from docx import Document

    doc = Document(ruta)
    txt = "\n".join(p.text for p in doc.paragraphs)
    assert "DOCUMENTO BORRADOR" in txt
    assert "t_arranque_ge_ms" in txt


def test_memoria_final_no_incluye_advertencia():
    datos = {
        **_datos_run_base(),
        "ats": {"t_arranque_ge_ms": 9000.0},
    }
    ruta = generar_memoria_sec(datos, _circuitos_base(), str(_tmp_dir()), modo_emision="auto")
    from docx import Document

    doc = Document(ruta)
    txt = "\n".join(p.text for p in doc.paragraphs)
    assert "DOCUMENTO BORRADOR" not in txt


def test_exportar_json_epc_incluye_nivel_incompleto_con_defaults():
    datos = {
        **_datos_run_base(),
        "ats": {"t_arranque_ge_ms": 10000.0},
    }
    ruta = exportar_json_epc(datos, str(_tmp_dir()))
    with open(ruta, "r", encoding="utf-8") as fh:
        payload = json.load(fh)
    assert payload["nivel_emision"] == "INCOMPLETO"
    assert payload["apto_emision"] is False
    assert payload["parametros_default"]


def test_exportar_json_epc_incluye_nivel_final_sin_defaults():
    datos = {
        **_datos_run_base(),
        "ats": {"t_arranque_ge_ms": 9000.0},
    }
    ruta = exportar_json_epc(datos, str(_tmp_dir()))
    with open(ruta, "r", encoding="utf-8") as fh:
        payload = json.load(fh)
    assert payload["nivel_emision"] == "FINAL"
    assert payload["apto_emision"] is True


# --- P0.3: seccion de alcance y supuestos ---

def test_memoria_incluye_seccion_alcance_supuestos():
    """P0.3: la memoria DOCX debe declarar alcance, supuestos y limitaciones."""
    from docx import Document as DocxDocument

    ruta = generar_memoria_docx(_datos_run_base(), _circuitos_base(), str(_tmp_dir()))
    doc = DocxDocument(ruta)
    textos = [p.text for p in doc.paragraphs]
    texto_completo = "\n".join(textos)

    assert "Alcance y Supuestos" in texto_completo
    assert "IEC 60909-2" in texto_completo
    assert "Limitaciones" in texto_completo


def test_memoria_alcance_menciona_impedancia_compleja():
    """La sección de alcance debe indicar el modelo R+jX."""
    from docx import Document as DocxDocument

    ruta = generar_memoria_docx(_datos_run_base(), _circuitos_base(), str(_tmp_dir()))
    doc = DocxDocument(ruta)
    texto_completo = "\n".join(p.text for p in doc.paragraphs)

    assert "R + jX" in texto_completo or "R+jX" in texto_completo


def test_memoria_alcance_menciona_rango_valido():
    """La sección de alcance debe declarar el rango de validez."""
    from docx import Document as DocxDocument

    ruta = generar_memoria_docx(_datos_run_base(), _circuitos_base(), str(_tmp_dir()))
    doc = DocxDocument(ruta)
    texto_completo = "\n".join(p.text for p in doc.paragraphs)

    assert "1000 V" in texto_completo or "1.5 mm" in texto_completo


def test_limitaciones_declara_sin_calculo_cuando_no_hay_datos():
    """Sin aporte de motores/cadena/proteccion, el disclaimer debe decir 'sin' en las 3."""
    from docx import Document as DocxDocument

    ruta = generar_memoria_docx(_datos_run_base(), _circuitos_base(), str(_tmp_dir()))
    doc = DocxDocument(ruta)
    texto_completo = "\n".join(p.text for p in doc.paragraphs)

    assert "Sin aporte de motores activado" in texto_completo
    assert "Sin análisis nodal de flujo de carga acoplado" in texto_completo
    assert "Sin Arc Flash calculado en esta memoria" in texto_completo


def test_limitaciones_no_declara_omision_cuando_hay_aporte_motores():
    """Si datos_run trae aporte_motores, el disclaimer no debe decir que está 'sin activar'."""
    from docx import Document as DocxDocument

    datos_run = dict(_datos_run_base())
    datos_run["aporte_motores"] = {"delta_icc_ka": 0.5}
    ruta = generar_memoria_docx(datos_run, _circuitos_base(), str(_tmp_dir()))
    doc = DocxDocument(ruta)
    texto_completo = "\n".join(p.text for p in doc.paragraphs)

    assert "Aporte de motores al Icc incluido en el cálculo de esta memoria" in texto_completo
    assert "Sin aporte de motores activado" not in texto_completo


def test_limitaciones_no_declara_omision_cuando_hay_cadena():
    """Si datos_run trae cadena (flujo nodal), el disclaimer no debe decir que está omitido."""
    from docx import Document as DocxDocument

    datos_run = dict(_datos_run_base())
    datos_run["cadena"] = [{"nombre": "C-01", "aguas_arriba": None}]
    ruta = generar_memoria_docx(datos_run, _circuitos_base(), str(_tmp_dir()))
    doc = DocxDocument(ruta)
    texto_completo = "\n".join(p.text for p in doc.paragraphs)

    assert "Análisis nodal de flujo de carga incluido en esta memoria" in texto_completo
    assert "Sin análisis nodal de flujo de carga acoplado" not in texto_completo


def test_limitaciones_no_declara_omision_cuando_hay_arc_flash_por_circuito():
    """Si algún circuito trae In_A/curva/icc_ka, el disclaimer debe reflejar Arc Flash calculado."""
    from docx import Document as DocxDocument

    circuitos = _circuitos_base()
    circuitos[0]["In_A"] = 40
    circuitos[0]["curva"] = "C"
    ruta = generar_memoria_docx(_datos_run_base(), circuitos, str(_tmp_dir()))
    doc = DocxDocument(ruta)
    texto_completo = "\n".join(p.text for p in doc.paragraphs)

    assert "Arc Flash IEEE 1584-2002 incluido en esta memoria" in texto_completo
    assert "Sin Arc Flash calculado en esta memoria" not in texto_completo
