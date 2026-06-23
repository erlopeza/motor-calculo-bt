"""P1.1 memoria — enriquecer circuitos y sección Arc Flash en la memoria SEC."""
from pathlib import Path
import uuid
from reporteria_sec import enriquecer_circuitos_con_proteccion


def test_enriquecer_agrega_in_y_curva():
    circuitos = [{"nombre": "C-01", "sistema": "3F", "icc_ka": 10.0}]
    protecciones = {"C-01": {"In_A": 250, "curva": "C", "poder_corte_kA": 25}}
    out = enriquecer_circuitos_con_proteccion(circuitos, protecciones)
    assert out[0]["In_A"] == 250
    assert out[0]["curva"] == "C"

def test_enriquecer_sin_proteccion_no_rompe():
    circuitos = [{"nombre": "C-99", "sistema": "3F", "icc_ka": 5.0}]
    out = enriquecer_circuitos_con_proteccion(circuitos, {})
    assert "In_A" not in out[0] or out[0].get("In_A") in (None, 0)

def test_enriquecer_no_muta_original():
    circuitos = [{"nombre": "C-01", "sistema": "3F", "icc_ka": 10.0}]
    enriquecer_circuitos_con_proteccion(circuitos, {"C-01": {"In_A": 100, "curva": "B"}})
    assert "In_A" not in circuitos[0]


def _tmp():
    d = Path(__file__).resolve().parent / ".tmp_arcflash"
    d.mkdir(parents=True, exist_ok=True)
    return str(d)

def _datos_run():
    return {
        "project_id": "TEST-AF", "revision": "R1",
        "timestamp": "2026-06-22T12:00:00+00:00", "perfil": "industrial",
        "norma": "MM2", "n_ok": 1, "n_fallas": 0, "max_dv_pct": 1.0,
        "max_icc_ka": 12.0, "status": "OK",
        "icc_barra_ka": 25.0,
        "proteccion_cabecera": {"In_A": 630, "curva": "C"},
    }

def _circuitos_af():
    return [
        {"nombre": "C-01", "conductor": "25mm2", "S_mm2": 25.0, "I_diseno": 80.0,
         "I_max": 100.0, "cos_phi": 0.9, "L_m": 20.0, "paralelos": 1, "sistema": "3F",
         "dv_v": 1.0, "dv_pct": 0.3, "icc_ka": 12.0, "estado": "OK", "norma": "MM2",
         "observaciones": "", "In_A": 100, "curva": "C"},
        {"nombre": "C-02-SIN-PROT", "conductor": "16mm2", "S_mm2": 16.0, "I_diseno": 50.0,
         "I_max": 76.0, "cos_phi": 0.9, "L_m": 30.0, "paralelos": 1, "sistema": "3F",
         "dv_v": 1.0, "dv_pct": 0.4, "icc_ka": 8.0, "estado": "OK", "norma": "MM2",
         "observaciones": ""},
    ]

def test_memoria_incluye_seccion_arc_flash():
    from docx import Document as DocxDocument
    from reporteria_sec import generar_memoria_docx
    ruta = generar_memoria_docx(_datos_run(), _circuitos_af(), _tmp())
    doc = DocxDocument(ruta)
    texto = "\n".join(p.text for p in doc.paragraphs)
    assert "Arco Eléctrico" in texto
    assert "IEEE 1584" in texto

def test_memoria_arc_flash_tiene_tabla_por_circuito():
    from docx import Document as DocxDocument
    from reporteria_sec import generar_memoria_docx
    ruta = generar_memoria_docx(_datos_run(), _circuitos_af(), _tmp())
    doc = DocxDocument(ruta)
    encontrada = any(
        any("Cat EPP" in cell.text for cell in tabla.rows[0].cells)
        for tabla in doc.tables
    )
    assert encontrada

def test_memoria_arc_flash_lista_circuitos_sin_proteccion():
    from docx import Document as DocxDocument
    from reporteria_sec import generar_memoria_docx
    ruta = generar_memoria_docx(_datos_run(), _circuitos_af(), _tmp())
    doc = DocxDocument(ruta)
    texto = "\n".join(p.text for p in doc.paragraphs)
    assert "C-02-SIN-PROT" in texto


def test_memoria_arc_flash_marca_despeje_incierto():
    """Un circuito cuya protección no despeja a Ia debe marcarse con ⚠ en la tabla."""
    from docx import Document as DocxDocument
    from reporteria_sec import generar_memoria_docx
    # icc baja + In enorme → no_dispara a Ia → despeje_incierto
    circuitos = [{
        "nombre": "C-INC", "conductor": "16mm2", "S_mm2": 16.0, "I_diseno": 40.0,
        "I_max": 76.0, "cos_phi": 0.9, "L_m": 30.0, "paralelos": 1, "sistema": "3F",
        "dv_v": 1.0, "dv_pct": 0.4, "icc_ka": 2.0, "estado": "OK", "norma": "MM2",
        "observaciones": "", "In_A": 4000, "curva": "C",
    }]
    ruta = generar_memoria_docx(_datos_run(), circuitos, _tmp())
    doc = DocxDocument(ruta)
    textos_tabla = [
        cell.text
        for tabla in doc.tables
        for row in tabla.rows
        for cell in row.cells
    ]
    assert any("⚠" in t for t in textos_tabla)
