from docx import Document

from src.generador_memoria import generar_memoria


def test_genera_archivo_docx(tmp_path):
    ruta = tmp_path / "memoria.docx"
    generar_memoria(datos_proyecto_fixture(), datos_calculo_fixture(), ruta)
    assert ruta.exists()
    assert ruta.suffix == ".docx"


def test_retorna_ok_true(tmp_path):
    resultado = generar_memoria(datos_proyecto_fixture(), datos_calculo_fixture(), tmp_path / "memoria.docx")
    assert resultado["ok"] is True


def test_archivo_no_vacio(tmp_path):
    ruta = tmp_path / "memoria.docx"
    generar_memoria(datos_proyecto_fixture(), datos_calculo_fixture(), ruta)
    assert ruta.stat().st_size > 0


def test_genera_con_emergencia_opcional(tmp_path):
    resultado = generar_memoria(datos_proyecto_fixture(), datos_calculo_con_opcionales(), tmp_path / "memoria.docx")
    assert resultado["ok"] is True


def test_genera_sin_emergencia(tmp_path):
    datos = datos_calculo_fixture()
    datos["emergencia"] = None
    resultado = generar_memoria(datos_proyecto_fixture(), datos, tmp_path / "memoria.docx")
    assert resultado["ok"] is True


def test_contiene_nombre_proyecto(tmp_path):
    ruta = tmp_path / "memoria.docx"
    generar_memoria(datos_proyecto_fixture(), datos_calculo_fixture(), ruta)
    assert "Proyecto SEC Demo" in _texto_doc(ruta)


def test_contiene_normativas_ric(tmp_path):
    ruta = tmp_path / "memoria.docx"
    generar_memoria(datos_proyecto_fixture(), datos_calculo_fixture(), ruta)
    texto = _texto_doc(ruta)
    assert "RIC N°03" in texto
    assert "RIC N°04" in texto


def test_tabla_circuitos_presente(tmp_path):
    ruta = tmp_path / "memoria.docx"
    generar_memoria(datos_proyecto_fixture(), datos_calculo_fixture(), ruta)
    assert Document(ruta).tables


def test_tabla_cubicacion_tiene_filas(tmp_path):
    ruta = tmp_path / "memoria.docx"
    generar_memoria(datos_proyecto_fixture(), datos_calculo_fixture(), ruta)
    tablas = Document(ruta).tables
    assert any(len(tabla.rows) > 2 for tabla in tablas)


def test_conclusion_contiene_valores_clave(tmp_path):
    ruta = tmp_path / "memoria.docx"
    generar_memoria(datos_proyecto_fixture(), datos_calculo_fixture(), ruta)
    assert "380" in _texto_doc(ruta)


def test_contiene_descripcion_obra(tmp_path):
    ruta = tmp_path / "memoria.docx"
    generar_memoria(datos_proyecto_fixture(), datos_calculo_fixture(), ruta)
    assert "Instalacion electrica de baja tension" in _texto_doc(ruta)


def test_contiene_puesta_tierra(tmp_path):
    ruta = tmp_path / "memoria.docx"
    generar_memoria(datos_proyecto_fixture(), datos_calculo_fixture(), ruta)
    assert "Sistema de Puesta a Tierra" in _texto_doc(ruta)


def test_contiene_conclusion_normativa(tmp_path):
    ruta = tmp_path / "memoria.docx"
    generar_memoria(datos_proyecto_fixture(), datos_calculo_fixture(), ruta)
    assert "Las instalaciones cumplen con normativa vigente." in _texto_doc(ruta)


def test_cubicacion_contiene_cable_agrupado(tmp_path):
    ruta = tmp_path / "memoria.docx"
    generar_memoria(datos_proyecto_fixture(), datos_calculo_fixture(), ruta)
    assert "Cable Cu 2.50 mm2" in _texto_tablas(ruta)


def test_cubicacion_contiene_interruptor(tmp_path):
    ruta = tmp_path / "memoria.docx"
    generar_memoria(datos_proyecto_fixture(), datos_calculo_fixture(), ruta)
    assert "Interruptor automatico 16A curva C" in _texto_tablas(ruta)


def test_arranque_opcional_aparece_en_texto(tmp_path):
    ruta = tmp_path / "memoria.docx"
    generar_memoria(datos_proyecto_fixture(), datos_calculo_con_opcionales(), ruta)
    assert "Arranque de motores" in _texto_doc(ruta)


def test_ruta_invalida_retorna_error(tmp_path):
    resultado = generar_memoria(datos_proyecto_fixture(), datos_calculo_fixture(), tmp_path / "no_existe" / "memoria.docx")
    assert resultado["ok"] is False


def test_datos_proyecto_incompleto_retorna_error(tmp_path):
    datos = datos_proyecto_fixture()
    datos.pop("nombre_proyecto")
    resultado = generar_memoria(datos, datos_calculo_fixture(), tmp_path / "memoria.docx")
    assert resultado["ok"] is False


def test_circuitos_vacios_retorna_error(tmp_path):
    datos = datos_calculo_fixture()
    datos["circuitos"] = []
    resultado = generar_memoria(datos_proyecto_fixture(), datos, tmp_path / "memoria.docx")
    assert resultado["ok"] is False


def datos_proyecto_fixture():
    return {
        "nombre_proyecto": "Proyecto SEC Demo",
        "direccion": "Av. Demo 123",
        "comuna": "Santiago",
        "propietario": "Cliente Demo",
        "rut_propietario": "11.111.111-1",
        "instalador": "Instalador Demo",
        "licencia": "12345",
        "clase_licencia": "A",
        "fecha": "2026-04",
        "resistencia_tierra_ohm": 12.5,
        "descripcion_obra": "Instalacion electrica de baja tension para proyecto demo.",
    }


def datos_calculo_fixture():
    return {
        "tension_v": 380.0,
        "potencia_total_kw": 12.5,
        "corriente_total_a": 24.0,
        "factor_potencia": 0.85,
        "circuitos": [
            {"nombre": "C1", "potencia_w": 2500, "corriente_a": 11.3, "seccion_mm2": 2.5, "ducto_mm": 20, "proteccion_a": 16, "curva": "C", "long_m": 20, "caida_v": 2.1, "caida_pct": 0.55},
            {"nombre": "C2", "potencia_w": 3500, "corriente_a": 15.9, "seccion_mm2": 4.0, "ducto_mm": 25, "proteccion_a": 20, "curva": "C", "long_m": 35, "caida_v": 3.0, "caida_pct": 0.78},
            {"nombre": "C3", "potencia_w": 1500, "corriente_a": 6.8, "seccion_mm2": 2.5, "ducto_mm": 20, "proteccion_a": 10, "curva": "C", "long_m": 12, "caida_v": 1.0, "caida_pct": 0.26},
        ],
        "alimentador": {"seccion_mm2": 10.0, "long_m": 50.0, "caida_v": 4.2, "caida_pct": 1.1, "tipo_cable": "Cu"},
        "emergencia": None,
        "arranque": None,
    }


def datos_calculo_con_opcionales():
    datos = datos_calculo_fixture()
    datos["emergencia"] = {"grupo": {"grupo": 1}, "autonomia": {"autonomia_min": 120}}
    datos["arranque"] = {"in_a": 9.71, "ia_arranque_a": 58.26}
    return datos


def _texto_doc(ruta):
    doc = Document(ruta)
    return "\n".join(parrafo.text for parrafo in doc.paragraphs)


def _texto_tablas(ruta):
    doc = Document(ruta)
    celdas = []
    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                celdas.append(celda.text)
    return "\n".join(celdas)
