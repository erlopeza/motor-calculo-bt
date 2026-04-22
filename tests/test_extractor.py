import uuid
from pathlib import Path

import pytest
from docx import Document
from reportlab.pdfgen import canvas

from rag_normativa.extractor import extraer_documento, extraer_docx, extraer_pdf


def _crear_pdf(path: Path, lineas=None):
    c = canvas.Canvas(str(path))
    y = 800
    for ln in lineas or []:
        c.drawString(72, y, ln)
        y -= 16
    c.save()


def _tmp_dir() -> Path:
    base = Path(__file__).resolve().parent / ".tmp_rag_extractor"
    base.mkdir(parents=True, exist_ok=True)
    d = base / uuid.uuid4().hex
    d.mkdir(parents=True, exist_ok=True)
    return d


def test_extraer_pdf_retorna_texto():
    ruta = _tmp_dir() / "norma.pdf"
    _crear_pdf(ruta, ["Articulo N°10", "Limite de caida de tension"])
    texto = extraer_pdf(str(ruta))
    assert "Articulo" in texto
    assert "caida de tension" in texto


def test_extraer_docx_retorna_texto():
    ruta = _tmp_dir() / "norma.docx"
    doc = Document()
    doc.add_paragraph("NCh 4/2003")
    doc.add_paragraph("Seccion minima conductores")
    doc.save(str(ruta))
    texto = extraer_docx(str(ruta))
    assert "NCh 4/2003" in texto
    assert "Seccion minima" in texto


def test_extraer_pdf_escaneado_retorna_vacio():
    ruta = _tmp_dir() / "scan.pdf"
    _crear_pdf(ruta, [])
    texto = extraer_pdf(str(ruta))
    assert texto == ""


def test_extraer_extension_no_soportada_raise():
    ruta = _tmp_dir() / "norma.xlsx"
    ruta.write_text("n/a", encoding="utf-8")
    with pytest.raises(ValueError):
        extraer_documento(str(ruta))
