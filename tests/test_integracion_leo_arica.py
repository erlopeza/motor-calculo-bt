"""Validación end-to-end sobre datos reales LEO-ARICA conformados al formato canónico."""
from pathlib import Path
import openpyxl
import pytest
from reporteria_sec import enriquecer_circuitos_con_proteccion, generar_memoria_docx

FIXTURE = Path(__file__).resolve().parent / "fixtures" / "leo_arica.xlsx"


def _leer_fixture():
    wb = openpyxl.load_workbook(FIXTURE, read_only=True, data_only=True)
    ws = wb["circuitos"]
    filas = list(ws.iter_rows(values_only=True))
    hdr = [str(h).lower() for h in filas[0]]
    circuitos, protecciones = [], {}
    for row in filas[1:]:
        d = dict(zip(hdr, row))
        nombre = d["nombre"]
        circuitos.append({
            "nombre": nombre, "sistema": d["sistema"], "conductor": d["conductor"],
            "S_mm2": float(str(d["conductor"]).replace("mm2", "")),
            "paralelos": int(d["paralelos"]), "I_diseno": float(d["i_diseno"]),
            "cos_phi": float(d["cos_phi"]), "L_m": float(d["l_m"]),
            "icc_ka": 15.0, "estado": "OK", "norma": "MM2",
        })
        protecciones[nombre] = {"In_A": d["in_a"], "curva": d["curva"]}
    return circuitos, protecciones


def test_fixture_existe():
    assert FIXTURE.exists(), "Ejecutar tests/_build_leo_fixture.py primero"

def test_memoria_leo_arica_genera_con_arc_flash(tmp_path):
    from docx import Document
    circuitos, protecciones = _leer_fixture()
    circuitos = enriquecer_circuitos_con_proteccion(circuitos, protecciones)
    datos_run = {
        "project_id": "LEO-ARICA", "revision": "VAL", "perfil": "datacenter",
        "norma": "MM2", "n_ok": len(circuitos), "n_fallas": 0,
        "max_dv_pct": 2.0, "max_icc_ka": 15.0, "status": "OK",
        "icc_barra_ka": 30.0, "proteccion_cabecera": {"In_A": 630, "curva": "C"},
    }
    ruta = generar_memoria_docx(datos_run, circuitos, str(tmp_path))
    doc = Document(ruta)
    texto = "\n".join(p.text for p in doc.paragraphs)
    assert "Arco Eléctrico" in texto
    tabla_af = [t for t in doc.tables if any("Cat EPP" in c.text for c in t.rows[0].cells)]
    assert tabla_af and len(tabla_af[0].rows) >= 2

def test_energia_en_rango_fisico(tmp_path):
    from arc_flash import arc_flash_desde_proteccion
    circuitos, protecciones = _leer_fixture()
    circuitos = enriquecer_circuitos_con_proteccion(circuitos, protecciones)
    for c in circuitos:
        r = arc_flash_desde_proteccion(15.0, 0.4, c["In_A"], c["curva"])
        assert 0 < r["E_cal_cm2"] < 1000, f"E fuera de rango en {c['nombre']}"
        assert r["D_afb_mm"] > 0
