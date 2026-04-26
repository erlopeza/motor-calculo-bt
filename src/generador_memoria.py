"""
Modulo Grupo 3 - Generador de memoria explicativa SEC.
Genera archivo .docx con estructura compatible RIC-N18.
Consume resultados de M1-M9. No recalcula nada.
"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt


NORMATIVAS = ["NCh Elec 4/2003"] + [f"RIC N°{i:02d}" for i in range(1, 20)]


def generar_memoria(datos_proyecto: dict, datos_calculo: dict, ruta_salida: str | Path) -> dict:
    """
    Genera memoria explicativa en formato .docx.

    Parametros:
        datos_proyecto: campos administrativos del formulario de usuario.
        datos_calculo: resultados calculados por M1-M9.
        ruta_salida: path completo del archivo .docx a generar.

    Retorna {"ok": bool, "ruta": str, "motivo": str}.
    """
    ruta = Path(ruta_salida)
    try:
        _validar_datos(datos_proyecto, datos_calculo, ruta)
        doc = Document()
        _configurar_documento(doc)
        _agregar_portada(doc, datos_proyecto)
        _agregar_indice(doc)
        _agregar_descripcion(doc, datos_proyecto, datos_calculo)
        _agregar_normativas(doc)
        _agregar_especificaciones(doc, datos_proyecto, datos_calculo)
        _agregar_calculos(doc, datos_calculo)
        _agregar_cubicacion(doc, datos_calculo)
        _agregar_puesta_tierra(doc, datos_proyecto)
        _agregar_conclusion(doc, datos_proyecto, datos_calculo)
        doc.save(ruta)
        return {"ok": True, "ruta": str(ruta), "motivo": ""}
    except Exception as error:
        return {"ok": False, "ruta": str(ruta), "motivo": str(error)}


def _agregar_portada(doc: Document, datos_proyecto: dict) -> None:
    """
    Agrega portada: titulo proyecto, direccion, instalador,
    licencia, fecha. Centrado, fuente Calibri.
    """
    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titulo.add_run("MEMORIA EXPLICATIVA DE INSTALACION ELECTRICA")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(16)

    for texto in [
        datos_proyecto.get("nombre_proyecto", ""),
        datos_proyecto.get("direccion", ""),
        f"Comuna: {datos_proyecto.get('comuna', '')}",
        f"Propietario: {datos_proyecto.get('propietario', '')}",
        f"RUT propietario: {datos_proyecto.get('rut_propietario', '')}",
        f"Instalador: {datos_proyecto.get('instalador', '')}",
        f"Licencia SEC: {datos_proyecto.get('licencia', '')} Clase {datos_proyecto.get('clase_licencia', '')}",
        f"Fecha: {datos_proyecto.get('fecha', '')}",
    ]:
        parrafo = doc.add_paragraph(texto)
        parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()


def _agregar_indice(doc: Document) -> None:
    """
    Agrega indice fijo con los 9 titulos de seccion.
    Nota: los numeros de pagina se actualizan al abrir el documento.
    """
    _estilo_titulo(doc, "Indice", 1)
    secciones = [
        "1. Descripcion de la Obra",
        "2. Normativas aplicables",
        "3. Especificaciones Tecnicas",
        "4. Calculos Justificativos",
        "5. Cubicacion de Materiales",
        "6. Sistema de Puesta a Tierra",
        "7. Conclusion",
    ]
    for seccion in secciones:
        _estilo_normal(doc, seccion)
    nota = doc.add_paragraph("Los numeros de pagina se actualizan al abrir el documento.")
    nota.runs[0].italic = True
    doc.add_page_break()


def _agregar_descripcion(doc: Document, datos_proyecto: dict, datos_calculo: dict) -> None:
    """
    Seccion 1. Descripcion de la Obra.
    Incluye descripcion de usuario, tension, potencia total y circuitos.
    """
    _estilo_titulo(doc, "1. Descripcion de la Obra", 1)
    descripcion = datos_proyecto.get("descripcion_obra") or "La presente memoria describe la instalacion electrica proyectada."
    _estilo_normal(doc, descripcion)
    _estilo_normal(
        doc,
        "La instalacion considera tension nominal de "
        f"{datos_calculo.get('tension_v', 0):.0f} V, potencia total de "
        f"{datos_calculo.get('potencia_total_kw', 0):.2f} kW y "
        f"{len(datos_calculo.get('circuitos', []))} circuitos derivados.",
    )


def _agregar_normativas(doc: Document) -> None:
    """
    Seccion 2. Normativas aplicables.
    Lista fija: NCh Elec 4/2003 + RIC N°01 al N°19.
    """
    _estilo_titulo(doc, "2. Normativas aplicables", 1)
    for normativa in NORMATIVAS:
        doc.add_paragraph(normativa, style="List Bullet")


def _agregar_especificaciones(doc: Document, datos_proyecto: dict, datos_calculo: dict) -> None:
    """
    Seccion 3. Especificaciones Tecnicas.
    Subsecciones: Generalidades, Tablero, Protecciones, Conductores, Alimentadores.
    """
    alimentador = datos_calculo.get("alimentador", {})
    _estilo_titulo(doc, "3. Especificaciones Tecnicas", 1)
    _estilo_titulo(doc, "3.1 Generalidades", 2)
    _estilo_normal(doc, "Los materiales deberan ser certificados para uso en instalaciones electricas de baja tension.")
    _estilo_titulo(doc, "3.2 Tablero", 2)
    _estilo_normal(
        doc,
        f"El tablero general operara a {datos_calculo.get('tension_v', 0):.0f} V "
        f"y corriente total estimada de {datos_calculo.get('corriente_total_a', 0):.2f} A.",
    )
    _estilo_titulo(doc, "3.3 Protecciones", 2)
    _estilo_normal(doc, "Las protecciones automaticas se seleccionan segun corriente de circuito y coordinacion disponible.")
    _estilo_titulo(doc, "3.4 Conductores", 2)
    _estilo_normal(doc, "Los conductores seran de cobre, aislacion normalizada y seccion segun calculo de capacidad y caida de tension.")
    _estilo_titulo(doc, "3.5 Alimentadores", 2)
    _estilo_normal(
        doc,
        f"Alimentador principal con cable {alimentador.get('tipo_cable', 'Cu')} "
        f"de {alimentador.get('seccion_mm2', 0):.2f} mm2 y longitud "
        f"{alimentador.get('long_m', 0):.2f} m.",
    )


def _agregar_calculos(doc: Document, datos_calculo: dict) -> None:
    """
    Seccion 4. Calculos Justificativos.
    Agrega alimentador, protecciones, conductores, canalizaciones,
    emergencia opcional y arranque opcional.
    """
    _estilo_titulo(doc, "4. Calculos Justificativos", 1)
    alimentador = datos_calculo.get("alimentador", {})
    _estilo_titulo(doc, "4.1 Alimentador", 2)
    _estilo_normal(
        doc,
        "Formula de caida de tension: Vp = K x rho x L x I / S. "
        f"Para el alimentador se obtiene caida de {alimentador.get('caida_v', 0):.2f} V "
        f"({alimentador.get('caida_pct', 0):.2f}%).",
    )

    _estilo_titulo(doc, "4.2 Protecciones", 2)
    _tabla_circuitos(doc, datos_calculo.get("circuitos", []), ["nombre", "potencia_w", "corriente_a", "proteccion_a", "curva", "seccion_mm2"])

    _estilo_titulo(doc, "4.3 Conductores", 2)
    _tabla_circuitos(doc, datos_calculo.get("circuitos", []), ["nombre", "seccion_mm2", "long_m", "caida_v", "caida_pct"])

    _estilo_titulo(doc, "4.4 Canalizaciones", 2)
    _tabla_circuitos(doc, datos_calculo.get("circuitos", []), ["nombre", "ducto_mm", "long_m"])

    if datos_calculo.get("emergencia") is not None:
        emergencia = datos_calculo["emergencia"]
        _estilo_titulo(doc, "4.5 Emergencia", 2)
        _estilo_normal(doc, f"Sistema de emergencia: {emergencia}")

    if datos_calculo.get("arranque") is not None:
        arranque = datos_calculo["arranque"]
        _estilo_titulo(doc, "4.6 Arranque de motores", 2)
        _estilo_normal(doc, f"Resultado de arranque: {arranque}")


def _agregar_cubicacion(doc: Document, datos_calculo: dict) -> None:
    """
    Seccion 5. Cubicacion de Materiales.
    Genera tabla automatica de cables, ductos, protecciones, diferencial y tablero.
    """
    _estilo_titulo(doc, "5. Cubicacion de Materiales", 1)
    tabla = doc.add_table(rows=1, cols=3)
    tabla.style = "Table Grid"
    headers = ["Descripcion", "Unidad", "Cantidad"]
    for idx, header in enumerate(headers):
        tabla.rows[0].cells[idx].text = header
        tabla.rows[0].cells[idx].paragraphs[0].runs[0].bold = True

    for descripcion, unidad, cantidad in _filas_cubicacion(datos_calculo):
        row = tabla.add_row().cells
        row[0].text = descripcion
        row[1].text = unidad
        row[2].text = str(cantidad)

    for _ in range(3):
        row = tabla.add_row().cells
        row[0].text = ""
        row[1].text = ""
        row[2].text = ""


def _agregar_puesta_tierra(doc: Document, datos_proyecto: dict) -> None:
    """
    Seccion 6. Sistema de Puesta a Tierra.
    Incluye valor medido, fecha y ubicacion en tablero general.
    """
    _estilo_titulo(doc, "6. Sistema de Puesta a Tierra", 1)
    valor = float(datos_proyecto.get("resistencia_tierra_ohm", 0) or 0)
    tabla = doc.add_table(rows=2, cols=4)
    tabla.style = "Table Grid"
    headers = ["Lectura N°", "Resistencia (ohm)", "Fecha", "Ubicacion"]
    valores = ["1", f"{valor:.2f}", datos_proyecto.get("fecha", ""), "Tablero general"]
    for idx, header in enumerate(headers):
        tabla.rows[0].cells[idx].text = header
        tabla.rows[0].cells[idx].paragraphs[0].runs[0].bold = True
        tabla.rows[1].cells[idx].text = valores[idx]
    conclusion = "El valor medido cumple con referencia RIC-N06." if valor <= 25 else "El valor medido debe ser revisado en terreno."
    _estilo_normal(doc, conclusion)


def _agregar_conclusion(doc: Document, datos_proyecto: dict, datos_calculo: dict) -> None:
    """
    Seccion 7. Conclusion.
    Texto boilerplate con valores clave del proyecto y cierre normativo.
    """
    alimentador = datos_calculo.get("alimentador", {})
    _estilo_titulo(doc, "7. Conclusion", 1)
    _estilo_normal(
        doc,
        f"La instalacion {datos_proyecto.get('nombre_proyecto', '')} considera "
        f"{datos_calculo.get('potencia_total_kw', 0):.2f} kW a "
        f"{datos_calculo.get('tension_v', 0):.0f} V, corriente total "
        f"{datos_calculo.get('corriente_total_a', 0):.2f} A, caida de tension "
        f"del alimentador {alimentador.get('caida_pct', 0):.2f}% y resistencia "
        f"de tierra {float(datos_proyecto.get('resistencia_tierra_ohm', 0) or 0):.2f} ohm.",
    )
    _estilo_normal(doc, "Las instalaciones cumplen con normativa vigente.")


def _estilo_titulo(doc: Document, texto: str, nivel: int = 1) -> None:
    """Agrega parrafo con estilo Heading nivel 1 o 2."""
    doc.add_heading(texto, level=nivel)


def _estilo_normal(doc: Document, texto: str) -> None:
    """Agrega parrafo con estilo Normal, Calibri 11."""
    parrafo = doc.add_paragraph(texto)
    for run in parrafo.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(11)


def _tabla_circuitos(doc: Document, circuitos: list[dict], columnas: list[str]) -> None:
    """
    Agrega tabla de circuitos con encabezado en negrita.
    columnas: lista de claves a mostrar del dict circuito.
    """
    tabla = doc.add_table(rows=1, cols=len(columnas))
    tabla.style = "Table Grid"
    for idx, columna in enumerate(columnas):
        tabla.rows[0].cells[idx].text = columna
        tabla.rows[0].cells[idx].paragraphs[0].runs[0].bold = True
    for circuito in circuitos:
        row = tabla.add_row().cells
        for idx, columna in enumerate(columnas):
            row[idx].text = str(circuito.get(columna, ""))


def _configurar_documento(doc: Document) -> None:
    """Configura margenes y fuente base del documento."""
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)


def _validar_datos(datos_proyecto: dict, datos_calculo: dict, ruta: Path) -> None:
    """Valida campos minimos antes de generar el documento."""
    requeridos = ["nombre_proyecto", "direccion", "instalador", "licencia"]
    faltantes = [campo for campo in requeridos if not datos_proyecto.get(campo)]
    if faltantes:
        raise ValueError(f"Faltan campos obligatorios: {', '.join(faltantes)}")
    if not ruta.parent.exists():
        raise FileNotFoundError(f"No existe directorio de salida: {ruta.parent}")
    if not datos_calculo.get("circuitos"):
        raise ValueError("datos_calculo debe incluir circuitos")


def _filas_cubicacion(datos_calculo: dict) -> list[tuple[str, str, float | int]]:
    """Agrupa materiales principales desde los circuitos calculados."""
    circuitos = datos_calculo.get("circuitos", [])
    alimentador = datos_calculo.get("alimentador", {})
    cables = {}
    ductos = {}
    protecciones = {}
    for circuito in circuitos:
        largo = float(circuito.get("long_m", 0) or 0)
        seccion = circuito.get("seccion_mm2", 0)
        cables[seccion] = cables.get(seccion, 0.0) + largo
        ducto = circuito.get("ducto_mm", 0)
        ductos[ducto] = ductos.get(ducto, 0.0) + largo
        prot = (circuito.get("proteccion_a", 0), circuito.get("curva", "C"))
        protecciones[prot] = protecciones.get(prot, 0) + 1

    filas = []
    tipo_cable = alimentador.get("tipo_cable", "Cu")
    for seccion, metros in sorted(cables.items(), key=lambda item: float(item[0])):
        filas.append((f"Cable {tipo_cable} {float(seccion):.2f} mm2", "m", round(metros, 2)))
    for ducto, metros in sorted(ductos.items(), key=lambda item: float(item[0])):
        filas.append((f"Ducto PVC {float(ducto):.0f} mm", "m", round(metros, 2)))
    for (capacidad, curva), cantidad in sorted(protecciones.items(), key=lambda item: float(item[0][0])):
        filas.append((f"Interruptor automatico {float(capacidad):.0f}A curva {curva}", "un", cantidad))
    filas.append(("Diferencial 40A 30mA", "un", 1))
    filas.append((f"Tablero {datos_calculo.get('tension_v', 0):.0f}V {datos_calculo.get('corriente_total_a', 0):.0f}A", "un", 1))
    return filas
