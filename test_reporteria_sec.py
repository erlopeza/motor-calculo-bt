import os
import re
import sqlite3
import uuid
from pathlib import Path

from persistencia import registrar_ejecucion
from reporteria_sec import (
    generar_desde_run_id,
    generar_memoria_docx,
    generar_memoria_sec,
    generar_reporte_pdf,
    verificar_completitud_parametros,
)


def _tmp_dir() -> Path:
    base = Path(__file__).resolve().parent / ".tmp_reporteria_sec_tests"
    base.mkdir(parents=True, exist_ok=True)
    return base


def _ruta_db_local(nombre: str) -> str:
    return str(_tmp_dir() / f"{nombre}_{uuid.uuid4().hex}.db")


def _datos_run_base():
    return {
        "project_id": "LEO-ARICA",
        "revision": "SEC1",
        "timestamp": "2026-04-19T12:00:00+00:00",
        "perfil": "industrial",
        "norma": "AWG",
        "n_ok": 2,
        "n_fallas": 0,
        "max_dv_pct": 1.5,
        "max_icc_ka": 20.0,
        "status": "OK",
    }


def _circuitos_base():
    return [
        {
            "nombre": "C-01",
            "conductor": "6AWG",
            "S_mm2": 13.3,
            "I_diseno": 42.0,
            "I_max": 65.0,
            "cos_phi": 0.9,
            "L_m": 15.0,
            "paralelos": 1,
            "sistema": "3F",
            "dv_v": 1.0,
            "dv_pct": 0.26,
            "icc_ka": 10.5,
            "estado": "OK",
            "norma": "AWG",
            "observaciones": "sin observaciones",
        }
    ]


def test_generar_memoria_docx():
    ruta = generar_memoria_docx(_datos_run_base(), _circuitos_base(), str(_tmp_dir()))
    assert ruta.endswith(".docx")
    assert Path(ruta).exists()


def test_generar_reporte_pdf():
    ruta = generar_reporte_pdf(_datos_run_base(), _circuitos_base(), str(_tmp_dir()))
    assert ruta.endswith(".pdf")
    assert Path(ruta).exists()


def test_nombres_archivo():
    ruta = generar_memoria_docx(_datos_run_base(), _circuitos_base(), str(_tmp_dir()))
    nombre = Path(ruta).name
    assert re.match(r"^MEMORIA_LEO-ARICA_SEC1_\d{8}_\d{4}\.docx$", nombre)


def test_generar_desde_run_id():
    ruta_db = _ruta_db_local("desde_run")
    run_id = registrar_ejecucion(
        {
            **_datos_run_base(),
            "ruta_reporte_txt": "reporte.txt",
            "ruta_reporte_xlsx": "reporte.xlsx",
            "circuitos": _circuitos_base(),
        },
        ruta_db=ruta_db,
    )

    cwd = os.getcwd()
    os.chdir(str(_tmp_dir()))
    try:
        ruta_docx, ruta_pdf = generar_desde_run_id(run_id, ruta_db=ruta_db)
    finally:
        os.chdir(cwd)

    assert ruta_docx.endswith(".docx")
    assert ruta_pdf.endswith(".pdf")
    assert Path(ruta_docx).exists()
    assert Path(ruta_pdf).exists()

    with sqlite3.connect(ruta_db) as conn:
        rows = conn.execute(
            "SELECT report_type, file_path FROM run_reports WHERE run_id = ?",
            (run_id,),
        ).fetchall()
    tipos = {r[0] for r in rows}
    assert "DOCX" in tipos
    assert "PDF" in tipos


def test_fallo_silencioso():
    ruta_db = _ruta_db_local("silencioso")
    try:
        rutas = generar_desde_run_id("run-id-no-existe", ruta_db=ruta_db)
    except Exception as e:
        raise AssertionError(f"No debia lanzar excepcion: {e}")
    assert rutas == ("", "")


def test_gate_emision_bloquea_con_defaults():
    gate = verificar_completitud_parametros(
        {
            "ats": {"t_arranque_ge_ms": 10000.0},
            "motor": {"factor_arranque": 6.0},
        }
    )
    assert gate["apto_emision"] is False
    assert gate["nivel"] == "INCOMPLETO"
    assert len(gate["parametros_default"]) >= 1


def test_gate_emision_aprueba_sin_defaults():
    gate = verificar_completitud_parametros(
        {
            "ats": {"t_arranque_ge_ms": 9000.0},
            "motor": {"factor_arranque": 5.5},
        }
    )
    assert gate["apto_emision"] is True
    assert gate["nivel"] == "FINAL"


def test_memoria_borrador_incluye_lista_defaults():
    datos = {
        **_datos_run_base(),
        "ats": {"t_arranque_ge_ms": 10000.0},
    }
    ruta = generar_memoria_sec(datos, _circuitos_base(), str(_tmp_dir()), modo_emision="auto")
    from docx import Document

    doc = Document(ruta)
    txt = "\n".join(p.text for p in doc.paragraphs)
    assert "DOCUMENTO BORRADOR" in txt
    assert "t_arranque_ge_ms" in txt


def test_memoria_final_no_incluye_advertencia():
    datos = {
        **_datos_run_base(),
        "ats": {"t_arranque_ge_ms": 9000.0},
    }
    ruta = generar_memoria_sec(datos, _circuitos_base(), str(_tmp_dir()), modo_emision="auto")
    from docx import Document

    doc = Document(ruta)
    txt = "\n".join(p.text for p in doc.paragraphs)
    assert "DOCUMENTO BORRADOR" not in txt
