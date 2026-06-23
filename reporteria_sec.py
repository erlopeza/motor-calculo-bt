"""Generacion de reporteria SEC desde resultados ya calculados."""

from __future__ import annotations

import os
import json
import sqlite3
import uuid
from datetime import datetime
from pathlib import Path
from typing import Tuple

from docx import Document
from docx.shared import Inches
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm
from reportlab.pdfgen import canvas

from persistencia import inicializar_db, obtener_circuitos, obtener_ejecuciones

PARAMETROS_DEFAULT_TIPO_A = [
    {
        "modulo": "ats",
        "variable": "t_arranque_ge_ms",
        "valor_default": 10000.0,
        "advertencia": "DEFAULT - verificar con ficha tecnica GE",
        "rutas": [("ats", "t_arranque_ge_ms"), ("t_arranque_ge_ms",)],
    },
    {
        "modulo": "ats",
        "variable": "t_estabilizacion_ms",
        "valor_default": 5000.0,
        "advertencia": "DEFAULT - verificar con ficha tecnica GE",
        "rutas": [("ats", "t_estabilizacion_ms"), ("t_estabilizacion_ms",)],
    },
    {
        "modulo": "ats",
        "variable": "t_deteccion_ms",
        "valor_default": 3000.0,
        "advertencia": "DEFAULT - configurable en AMF25",
        "rutas": [("ats", "t_deteccion_ms"), ("t_deteccion_ms",)],
    },
    {
        "modulo": "motores",
        "variable": "factor_arranque",
        "valor_default": 6.0,
        "advertencia": "DEFAULT - DOL tipico, verificar con motor real",
        "rutas": [("motor", "factor_arranque"), ("factor_arranque",)],
    },
    {
        "modulo": "ups",
        "variable": "prof_descarga_pct",
        "valor_default": 80.0,
        "advertencia": "DEFAULT - AGM, verificar con banco real",
        "rutas": [("ups", "prof_descarga_pct"), ("prof_descarga_pct",)],
    },
]


def _sanitize(value) -> str:
    if value is None:
        return "NA"
    return str(value).strip().replace(" ", "_")


def _marca_archivo() -> str:
    return datetime.now().strftime("%Y%m%d_%H%M")


def _asegurar_carpeta(ruta_salida: str) -> Path:
    carpeta = Path(ruta_salida).resolve()
    carpeta.mkdir(parents=True, exist_ok=True)
    return carpeta


def _ruta_memoria(datos_run: dict, ruta_salida: str) -> Path:
    carpeta = _asegurar_carpeta(ruta_salida)
    nombre = (
        f"MEMORIA_{_sanitize(datos_run.get('project_id'))}_"
        f"{_sanitize(datos_run.get('revision'))}_{_marca_archivo()}.docx"
    )
    return carpeta / nombre


def _ruta_reporte(datos_run: dict, ruta_salida: str) -> Path:
    carpeta = _asegurar_carpeta(ruta_salida)
    nombre = (
        f"REPORTE_{_sanitize(datos_run.get('project_id'))}_"
        f"{_sanitize(datos_run.get('revision'))}_{_marca_archivo()}.pdf"
    )
    return carpeta / nombre


def _fecha_formateada(datos_run: dict) -> str:
    raw = datos_run.get("timestamp", "")
    try:
        return datetime.fromisoformat(str(raw)).strftime("%d/%m/%Y %H:%M")
    except Exception:
        return str(raw) or "Sin fecha"


def _norma_display(datos_run: dict) -> str:
    return (
        "RIC N°04 / IEC 60228 — Conductores AWG"
        if datos_run.get("norma") == "AWG"
        else "RIC N°04 / IEC 60228 — Conductores MM2"
    )


def _potencia_activa(circuito: dict) -> float:
    i_dis = circuito.get("I_diseno") or 0
    cos_phi = circuito.get("cos_phi") or 0
    return round(float(i_dis) * 380.0 * float(cos_phi), 2)


def _valor_icc(icc) -> float:
    try:
        return float(icc or 0)
    except Exception:
        return 0.0


def _hay_icc_declarado(circuitos: list) -> bool:
    return any(_valor_icc(c.get("icc_ka")) > 0 for c in circuitos)


def _circuitos_en_falla(circuitos: list) -> list:
    en_falla = []
    for c in circuitos:
        estado = str(c.get("estado") or "").upper()
        dv_pct = c.get("dv_pct")
        try:
            dv_val = float(dv_pct)
        except Exception:
            dv_val = 0.0
        if "FALLA" in estado or dv_val > 5.0:
            en_falla.append(str(c.get("nombre") or "SIN_NOMBRE"))
    return en_falla


def _criterio_normativo(modulo: str, parametro: str) -> str:
    try:
        from rag_normativa import consultar_criterio_calculo

        criterio = consultar_criterio_calculo(modulo, parametro)
        if criterio.get("ok"):
            return str(criterio.get("cita") or "")
    except Exception:
        return ""
    return ""


def _valor_por_ruta(base: dict, ruta: tuple) -> object:
    cur = base
    for key in ruta:
        if not isinstance(cur, dict):
            return None
        cur = cur.get(key)
    return cur


def verificar_completitud_parametros(resultados: dict) -> dict:
    """
    Verifica que parametros TIPO A no queden en DEFAULT sin confirmacion.
    """
    confirmados = set(resultados.get("defaults_confirmados") or [])
    parametros_default = []
    for spec in PARAMETROS_DEFAULT_TIPO_A:
        val = None
        for ruta in spec["rutas"]:
            got = _valor_por_ruta(resultados, ruta)
            if got is not None:
                val = got
                break
        if val is None:
            continue
        try:
            if abs(float(val) - float(spec["valor_default"])) < 1e-9:
                parametros_default.append(
                    {
                        "modulo": spec["modulo"],
                        "variable": spec["variable"],
                        "valor": val,
                        "advertencia": spec["advertencia"],
                        "confirmado": spec["variable"] in confirmados,
                    }
                )
        except Exception:
            continue

    if not parametros_default:
        return {"apto_emision": True, "parametros_default": [], "nivel": "FINAL"}

    todos_confirmados = all(p["confirmado"] for p in parametros_default)
    if todos_confirmados:
        return {"apto_emision": True, "parametros_default": parametros_default, "nivel": "BORRADOR"}
    return {"apto_emision": False, "parametros_default": parametros_default, "nivel": "INCOMPLETO"}


def _resultados_para_graficos(datos_run: dict, circuitos: list) -> dict:
    max_icc = datos_run.get("max_icc_ka")
    return {
        "circuitos": [
            {"id": c.get("nombre") or f"C-{idx + 1:02d}", "dv_pct": float(c.get("dv_pct") or 0.0)}
            for idx, c in enumerate(circuitos)
        ],
        "generador": (datos_run.get("generador") or None),
        "protecciones": (datos_run.get("protecciones") or None),
        "Icc_punto_kA": None if max_icc is None else float(max_icc),
        "balance": (datos_run.get("balance") or None),
        "ups": (datos_run.get("ups") or None),
        "ats": (datos_run.get("ats") or None),
        "simulaciones": (datos_run.get("simulaciones") or None),
        "commissioning": (datos_run.get("commissioning") or None),
    }


def _insertar_grafico_docx(doc: Document, titulo: str, ruta_imagen: str) -> None:
    if not ruta_imagen:
        return
    if not Path(ruta_imagen).exists():
        return
    doc.add_paragraph(titulo)
    doc.add_picture(ruta_imagen, width=Inches(6.5))


def _agregar_seccion_alcance_supuestos(doc: Document) -> None:
    """Sección fija P0.3: alcance, supuestos y limitaciones del modelo de cálculo."""
    doc.add_heading("Alcance y Supuestos del Modelo de Cálculo", level=1)

    doc.add_heading("Ámbito de aplicación", level=2)
    doc.add_paragraph(
        "El presente cálculo aplica a instalaciones de baja tensión (BT) "
        "hasta 1000 V, 50 Hz, según normativa chilena SEC/NCh e internacional IEC."
    )

    doc.add_heading("Modelo de impedancia de cable", level=2)
    doc.add_paragraph(
        "Impedancia compleja Z = R + jX por tramo de conductor:\n"
        "  R = ρ × L / (S × n_paralelos)  [IEC 60228, ρ_Cu = 0.0175 Ω·mm²/m]\n"
        "  X interpolado de tabla IEC 60909-2:2008 Tabla B.1 (50 Hz, Cu, conduit/bandeja)"
    )

    doc.add_heading("Cortocircuito", level=2)
    doc.add_paragraph(
        "Método: IEC 60909 con factores c_max = 1.05 / c_min = 0.95.\n"
        "Se calculan Icc trifásica (3F) e Icc fase-neutro (IEC 60364-4-41) en cada punto.\n"
        "Aporte de motores (IEC 60909-4:2021): disponible por llamada explícita a "
        "calcular_icc_con_aporte_motores(); no incluido por defecto en esta memoria."
    )

    doc.add_heading("Rango de validez", level=2)
    doc.add_paragraph(
        "• Tensión de sistema: 220 V (1F/2F) o 380 V (3F) fase-línea.\n"
        "• Sección de conductor: 1.5 mm² a 500 mm² (Cu) o equivalente AWG.\n"
        "• Longitud de tramo: sin límite superior, con X ganando peso relativo > 50 m.\n"
        "• Temperatura de referencia para R: 20°C (IEC 60228); corrección por temperatura\n"
        "  disponible en calculos.py (NEC Table 310.15(B)(1))."
    )

    doc.add_heading("Limitaciones conocidas", level=2)
    doc.add_paragraph(
        "• Impedancia de transformador: tratada como puramente resistiva en esta versión\n"
        "  (X_trafo no incluida; error < 2% en BT típico).\n"
        "• Sin aporte de motores activado en el cálculo base de esta memoria.\n"
        "• Sin análisis nodal de flujo de carga acoplado (calculista por circuito).\n"
        "• Sin Arc Flash IEEE 1584-2018 (previsto en fase F2 del roadmap)."
    )


def _agregar_seccion_arc_flash(doc: Document, datos_run: dict, circuitos: list) -> None:
    """Seccion de Arc Flash (IEEE 1584): barra principal + tabla por circuito."""
    from arc_flash import arc_flash_desde_proteccion
    from conductores import TENSION_SISTEMA

    doc.add_heading("Análisis de Arco Eléctrico (IEEE 1584)", level=1)

    # --- Barra principal ---
    icc_barra = datos_run.get("icc_barra_ka")
    cab = datos_run.get("proteccion_cabecera") or {}
    # Tensión de barra coherente con la usada por circuito (TENSION_SISTEMA, en kV).
    v_barra_kv = float(datos_run.get("tension_barra_kv") or TENSION_SISTEMA.get("3F", 380) / 1000.0)
    if icc_barra and cab.get("In_A") and cab.get("curva"):
        r = arc_flash_desde_proteccion(
            float(icc_barra), v_barra_kv, cab["In_A"], cab["curva"]
        )
        cat = "PELIGRO" if r["categoria_ppe"] is None else f"Cat {r['categoria_ppe']}"
        doc.add_heading("Barra principal", level=2)
        doc.add_paragraph(
            f"Ibf={r['Ibf_kA']:.2f} kA - Ia={r['Ia_kA']:.2f} kA - "
            f"t_despeje={r['t_despeje_s']:.3f} s - E={r['E_cal_cm2']:.2f} cal/cm2 - "
            f"Frontera={r['D_afb_mm']:.0f} mm - {cat}"
            + ("  ⚠ despeje incierto" if r["despeje_incierto"] else "")
        )
    else:
        doc.add_paragraph(
            "Barra principal: sin Icc de barra o proteccion de cabecera definida; "
            "se omite el calculo de arco en la barra."
        )

    # --- Tabla por circuito ---
    con_prot = [c for c in circuitos if c.get("In_A") and c.get("curva") and c.get("icc_ka")]
    # Complemento exacto de con_prot: incluye también los que tienen protección
    # pero les falta icc_ka (no quedan en el limbo, se reportan como omitidos).
    sin_prot = [c for c in circuitos if not (c.get("In_A") and c.get("curva") and c.get("icc_ka"))]

    if con_prot:
        doc.add_heading("Por circuito", level=2)
        tabla = doc.add_table(rows=1, cols=7)
        encabezados = ["Circuito", "Icc (kA)", "Ia (kA)", "t_desp (s)",
                       "E (cal/cm2)", "Frontera (mm)", "Cat EPP"]
        for i, h in enumerate(encabezados):
            tabla.rows[0].cells[i].text = h
        for c in con_prot:
            v_kv = TENSION_SISTEMA.get(c["sistema"], 380) / 1000.0
            r = arc_flash_desde_proteccion(
                float(c["icc_ka"]), v_kv, c["In_A"], c["curva"]
            )
            marca = " ⚠" if (r["despeje_incierto"] or r["verificar_simaris"]) else ""
            cat = "PELIGRO" if r["categoria_ppe"] is None else str(r["categoria_ppe"])
            celdas = tabla.add_row().cells
            celdas[0].text = str(c.get("nombre", ""))
            celdas[1].text = f'{float(c["icc_ka"]):.2f}'
            celdas[2].text = f'{r["Ia_kA"]:.2f}'
            celdas[3].text = f'{r["t_despeje_s"]:.3f}'
            celdas[4].text = f'{r["E_cal_cm2"]:.2f}'
            celdas[5].text = f'{r["D_afb_mm"]:.0f}'
            celdas[6].text = cat + marca

    if sin_prot:
        nombres = ", ".join(str(c.get("nombre", "?")) for c in sin_prot)
        doc.add_paragraph(
            "Circuitos sin datos de proteccion (omitidos del analisis de arco): "
            + nombres
        )


def enriquecer_circuitos_con_proteccion(circuitos: list, protecciones: dict) -> list:
    """Devuelve copias de los circuitos con In_A/curva de su protección (por nombre).

    `protecciones` es {nombre_circuito: {"In_A": .., "curva": .., ...}}.
    Circuitos sin protección quedan sin In_A/curva. No muta la entrada.
    """
    salida = []
    for c in circuitos:
        c2 = dict(c)
        prot = protecciones.get(c.get("nombre"))
        if prot:
            if prot.get("In_A"):
                c2["In_A"] = prot["In_A"]
            if prot.get("curva"):
                c2["curva"] = prot["curva"]
        salida.append(c2)
    return salida


def generar_memoria_docx(
    datos_run: dict,
    circuitos: list,
    ruta_salida: str
) -> str:
    ruta_docx = _ruta_memoria(datos_run, ruta_salida)
    fecha = _fecha_formateada(datos_run)
    graficos_paths = {}
    try:
        from graficos import generar_todos as _gen_graficos

        ruta_curvas = str(_asegurar_carpeta(os.path.join(ruta_salida, "07_CONTROL", "curvas")))
        graficos_paths = _gen_graficos(
            _resultados_para_graficos(datos_run, circuitos),
            ruta_curvas,
            prefijo=f"{_sanitize(datos_run.get('project_id'))}_",
        )
    except Exception:
        graficos_paths = {}

    doc = Document()
    titulo = (
        "MEMORIA EXPLICATIVA DE CÁLCULO\n"
        f"Instalación Eléctrica Interior — {datos_run.get('project_id')}\n"
        f"Revisión: {datos_run.get('revision')}"
    )
    doc.add_heading(titulo, level=0)

    doc.add_heading("Portada", level=1)
    doc.add_paragraph(f"Proyecto: {datos_run.get('project_id')}")
    doc.add_paragraph(f"Revision: {datos_run.get('revision')}")
    doc.add_paragraph(f"Fecha: {fecha}")
    doc.add_paragraph(f"Norma: {datos_run.get('norma')}")
    doc.add_paragraph(f"Perfil: {datos_run.get('perfil')}")
    parametros_default_sec = datos_run.get("_parametros_default_sec") or []
    nivel_emision = str(datos_run.get("_nivel_emision_sec") or "").upper()
    if parametros_default_sec and nivel_emision in {"BORRADOR", "INCOMPLETO"}:
        doc.add_paragraph("⚠ DOCUMENTO BORRADOR")
        doc.add_paragraph("Los siguientes parámetros usan valores por defecto")
        doc.add_paragraph("pendientes de verificación con ficha técnica:")
        for p in parametros_default_sec:
            doc.add_paragraph(
                f"  · {p.get('variable')} = {p.get('valor')} ({p.get('advertencia')})"
            )
        doc.add_paragraph("Confirme estos valores antes de emitir versión final.")

    doc.add_heading("Antecedentes", level=1)
    doc.add_paragraph("Sistema de referencia: 380V / 3F / 50Hz")
    doc.add_paragraph(f"Norma aplicada: {_norma_display(datos_run)}")
    cita_dv = _criterio_normativo("calculos", "limite_dv")
    if cita_dv:
        doc.add_paragraph(f"Criterio aplicado: {cita_dv}")

    _agregar_seccion_alcance_supuestos(doc)

    transformador = datos_run.get("transformador") or {}
    if transformador:
        doc.add_heading("Transformador", level=1)
        doc.add_paragraph(f"Potencia nominal: {transformador.get('kVA')} kVA")
        doc.add_paragraph(f"Tension BT: {transformador.get('Vn_BT')} V")
        doc.add_paragraph(f"Ucc%: {transformador.get('Ucc_pct')} %")
        doc.add_paragraph(
            "Icc bornes BT: "
            f"nominal={transformador.get('Icc_nom_kA')} kA, "
            f"maxima={transformador.get('Icc_max_kA')} kA, "
            f"minima={transformador.get('Icc_min_kA')} kA"
        )

    doc.add_heading("Cuadro de cargas", level=1)
    tabla_cargas = doc.add_table(rows=1, cols=4)
    tabla_cargas.rows[0].cells[0].text = "Circuito"
    tabla_cargas.rows[0].cells[1].text = "Conductor"
    tabla_cargas.rows[0].cells[2].text = "I_diseno (A)"
    tabla_cargas.rows[0].cells[3].text = "Potencia Activa (W)"
    for c in circuitos:
        row = tabla_cargas.add_row().cells
        row[0].text = str(c.get("nombre", ""))
        row[1].text = str(c.get("conductor", ""))
        row[2].text = str(c.get("I_diseno") or 0)
        row[3].text = str(_potencia_activa(c))

    doc.add_heading("Calculo DeltaV", level=1)
    tabla_dv = doc.add_table(rows=1, cols=6)
    tabla_dv.rows[0].cells[0].text = "Circuito"
    tabla_dv.rows[0].cells[1].text = "L_m"
    tabla_dv.rows[0].cells[2].text = "S_mm2"
    tabla_dv.rows[0].cells[3].text = "dv_v"
    tabla_dv.rows[0].cells[4].text = "dv_pct"
    tabla_dv.rows[0].cells[5].text = "estado"
    for c in circuitos:
        row = tabla_dv.add_row().cells
        row[0].text = str(c.get("nombre", ""))
        row[1].text = str(c.get("L_m", ""))
        row[2].text = str(c.get("S_mm2", ""))
        row[3].text = str(c.get("dv_v", ""))
        row[4].text = str(c.get("dv_pct", ""))
        row[5].text = str(c.get("estado", ""))

    if _hay_icc_declarado(circuitos):
        doc.add_heading("Calculo Icc", level=1)
        tabla_icc = doc.add_table(rows=1, cols=4)
        tabla_icc.rows[0].cells[0].text = "Circuito"
        tabla_icc.rows[0].cells[1].text = "icc_ka"
        tabla_icc.rows[0].cells[2].text = "I_max"
        tabla_icc.rows[0].cells[3].text = "estado proteccion"
        for c in circuitos:
            row = tabla_icc.add_row().cells
            row[0].text = str(c.get("nombre", ""))
            row[1].text = str(c.get("icc_ka", ""))
            row[2].text = str(c.get("I_max", ""))
            row[3].text = str(c.get("estado", ""))
    else:
        doc.add_paragraph(
            "Cálculo de cortocircuito: no declarado en este proyecto.\n"
            "Verificar con herramienta SIMARIS si aplica."
        )

    try:
        _agregar_seccion_arc_flash(doc, datos_run, circuitos)
    except Exception as e:
        print(f"[reporteria_sec] Arc Flash no disponible: {e}")
        doc.add_paragraph(f"[Análisis de Arco Eléctrico no disponible: {e}]")

    balance_demanda = datos_run.get("balance_demanda") or {}
    if balance_demanda:
        doc.add_heading("Balance y Demanda", level=1)
        balance = balance_demanda.get("balance") or {}
        demanda = balance_demanda.get("demanda") or {}
        if balance:
            doc.add_paragraph(
                "Balance: "
                f"Demanda total={balance.get('S_total_kva')} kVA, "
                f"Uso transformador={balance.get('uso_trafo_pct')}%"
            )
        if demanda:
            doc.add_paragraph(
                "Demanda: "
                f"{demanda.get('S_total_kva')} kVA / {demanda.get('P_total_kw')} kW, "
                f"Factor crecimiento={demanda.get('factor_crecimiento')}, "
                f"Demanda futura={demanda.get('S_futuro_kva')} kVA"
            )

    if graficos_paths:
        doc.add_heading("Graficos tecnicos", level=1)
        _insertar_grafico_docx(doc, "Perfil DeltaV por circuito", graficos_paths.get("dv_circuitos"))
        _insertar_grafico_docx(doc, "Curva de decremento Icc GE", graficos_paths.get("decremento_ge"))
        _insertar_grafico_docx(doc, "Curvas TCC", graficos_paths.get("tcc"))
        _insertar_grafico_docx(doc, "Balance de carga por fase", graficos_paths.get("balance_fases"))
        _insertar_grafico_docx(doc, "Autonomia UPS", graficos_paths.get("autonomia_ups"))
        _insertar_grafico_docx(doc, "Secuencia ATS", graficos_paths.get("transferencia_ats"))
        _insertar_grafico_docx(doc, "Divergencias vs SIMARIS", graficos_paths.get("divergencias_simaris"))
        _insertar_grafico_docx(doc, "Estado commissioning", graficos_paths.get("commissioning"))

    doc.add_heading("Conclusion", level=1)
    icc_str = (
        f"{datos_run['max_icc_ka']} kA"
        if datos_run.get("max_icc_ka") and datos_run["max_icc_ka"] > 0
        else "No calculado"
    )
    doc.add_paragraph(f"n_ok: {datos_run.get('n_ok')}")
    doc.add_paragraph(f"n_fallas: {datos_run.get('n_fallas')}")
    doc.add_paragraph(f"max_dv_pct: {datos_run.get('max_dv_pct')}")
    doc.add_paragraph(f"max_icc_ka: {icc_str}")
    doc.add_paragraph(f"status: {datos_run.get('status')}")

    n_fallas = int(datos_run.get("n_fallas") or 0)
    if n_fallas > 0:
        doc.add_paragraph(
            f"ATENCIÓN: {n_fallas} circuito(s) presentan caída de tensión superior\n"
            "al límite normativo (ΔV > 5%, RIC N°10).\n"
            "Acción requerida: redimensionar conductor antes de declarar\n"
            "ante SEC."
        )
        en_falla = _circuitos_en_falla(circuitos)
        if en_falla:
            doc.add_paragraph("Circuitos en falla: " + ", ".join(en_falla))

    doc.save(str(ruta_docx))
    return str(ruta_docx)


def generar_memoria_sec(
    datos_run: dict,
    circuitos: list,
    ruta_salida: str,
    modo_emision: str = "auto",
) -> str:
    """
    Genera memoria SEC con gate de completitud de parametros.
    """
    modo = str(modo_emision or "auto").lower()
    gate = verificar_completitud_parametros(datos_run)
    if modo == "final":
        gate = {"apto_emision": True, "parametros_default": [], "nivel": "FINAL"}
    elif modo == "borrador":
        gate = {"apto_emision": True, "parametros_default": gate["parametros_default"], "nivel": "BORRADOR"}

    datos_memoria = dict(datos_run)
    datos_memoria["_parametros_default_sec"] = gate.get("parametros_default") or []
    datos_memoria["_nivel_emision_sec"] = gate.get("nivel")
    return generar_memoria_docx(datos_memoria, circuitos, ruta_salida)


def exportar_json_epc(
    datos_run: dict,
    ruta_salida: str,
    modo_emision: str = "auto",
) -> str:
    """
    Exporta un JSON minimo para consumo EPC con nivel de emision SEC.
    """
    modo = str(modo_emision or "auto").lower()
    gate = verificar_completitud_parametros(datos_run)
    if modo == "final":
        gate = {"apto_emision": True, "parametros_default": [], "nivel": "FINAL"}
    elif modo == "borrador":
        gate = {"apto_emision": True, "parametros_default": gate["parametros_default"], "nivel": "BORRADOR"}

    carpeta = _asegurar_carpeta(ruta_salida)
    nombre = (
        f"EPC_{_sanitize(datos_run.get('project_id'))}_"
        f"{_sanitize(datos_run.get('revision'))}_{_marca_archivo()}.json"
    )
    ruta_json = carpeta / nombre
    payload = {
        "project_id": datos_run.get("project_id"),
        "revision": datos_run.get("revision"),
        "timestamp": datos_run.get("timestamp"),
        "nivel_emision": gate.get("nivel"),
        "apto_emision": gate.get("apto_emision"),
        "parametros_default": gate.get("parametros_default") or [],
        "resultados": datos_run,
    }
    with open(ruta_json, "w", encoding="utf-8") as fh:
        json.dump(payload, fh, ensure_ascii=False, indent=2, default=str)
    return str(ruta_json)


def generar_reporte_pdf(
    datos_run: dict,
    circuitos: list,
    ruta_salida: str
) -> str:
    ruta_pdf = _ruta_reporte(datos_run, ruta_salida)
    fecha = _fecha_formateada(datos_run)

    c = canvas.Canvas(str(ruta_pdf), pagesize=A4)
    _, h = A4

    y = h - 2 * cm
    c.setFont("Helvetica-Bold", 14)
    c.drawString(2 * cm, y, "MEMORIA EXPLICATIVA DE CÁLCULO")
    y -= 0.7 * cm
    c.drawString(2 * cm, y, f"Instalación Eléctrica Interior — {datos_run.get('project_id')}")
    y -= 0.7 * cm
    c.drawString(2 * cm, y, f"Revisión: {datos_run.get('revision')}")

    y -= 0.8 * cm
    c.setFont("Helvetica", 10)
    c.drawString(2 * cm, y, f"Fecha: {fecha}")
    y -= 0.5 * cm
    c.drawString(2 * cm, y, f"Norma aplicada: {_norma_display(datos_run)}")

    y -= 1.0 * cm
    c.setFont("Helvetica-Bold", 10)
    c.drawString(2 * cm, y, "Cuadro de cargas resumido")

    y -= 0.7 * cm
    c.setFont("Helvetica", 9)
    c.drawString(2 * cm, y, "Circuito")
    c.drawString(7.8 * cm, y, "Conductor")
    c.drawString(11.0 * cm, y, "Potencia Activa (W)")
    c.drawString(15.0 * cm, y, "dv_pct")
    c.drawString(17.0 * cm, y, "Estado")

    y -= 0.3 * cm
    c.line(2 * cm, y, 19 * cm, y)

    for circuito in circuitos:
        y -= 0.5 * cm
        if y < 2 * cm:
            c.showPage()
            y = h - 2 * cm
            c.setFont("Helvetica", 9)
        c.drawString(2 * cm, y, str(circuito.get("nombre", ""))[:28])
        c.drawString(7.8 * cm, y, str(circuito.get("conductor", ""))[:15])
        c.drawString(11.0 * cm, y, str(_potencia_activa(circuito))[:16])
        c.drawString(15.0 * cm, y, str(circuito.get("dv_pct", ""))[:8])
        c.drawString(17.0 * cm, y, str(circuito.get("estado", ""))[:12])

    y -= 1.0 * cm
    if y < 3 * cm:
        c.showPage()
        y = h - 2 * cm
    c.setFont("Helvetica-Bold", 10)
    c.drawString(2 * cm, y, "Conclusion")
    y -= 0.6 * cm
    c.setFont("Helvetica", 10)

    icc_str = (
        f"{datos_run['max_icc_ka']} kA"
        if datos_run.get("max_icc_ka") and datos_run["max_icc_ka"] > 0
        else "No calculado"
    )
    c.drawString(2 * cm, y, f"Status final: {datos_run.get('status')}")
    y -= 0.5 * cm
    c.drawString(2 * cm, y, f"Circuitos OK: {datos_run.get('n_ok')} | Fallas: {datos_run.get('n_fallas')}")
    y -= 0.5 * cm
    c.drawString(2 * cm, y, f"Icc máximo: {icc_str}")

    n_fallas = int(datos_run.get("n_fallas") or 0)
    if n_fallas > 0:
        y -= 0.6 * cm
        c.drawString(2 * cm, y, f"ATENCIÓN: {n_fallas} circuito(s) superan ΔV > 5% (RIC N°10).")
        y -= 0.5 * cm
        c.drawString(2 * cm, y, "Acción requerida: redimensionar conductor antes de declarar ante SEC.")
        en_falla = _circuitos_en_falla(circuitos)
        if en_falla:
            y -= 0.5 * cm
            c.drawString(2 * cm, y, "Circuitos en falla: " + ", ".join(en_falla)[:110])

    c.save()
    return str(ruta_pdf)


def _registrar_ruta_reporte(run_id: str, report_type: str, file_path: str, ruta_db: str) -> None:
    try:
        inicializar_db(ruta_db)
        with sqlite3.connect(ruta_db) as conn:
            conn.execute("PRAGMA foreign_keys = ON")
            conn.execute(
                """
                INSERT INTO run_reports (report_id, run_id, report_type, file_path)
                VALUES (?, ?, ?, ?)
                ON CONFLICT(run_id, report_type)
                DO UPDATE SET file_path=excluded.file_path
                """,
                (str(uuid.uuid4()), run_id, report_type, file_path),
            )
            conn.commit()
    except Exception:
        print("[reporteria_sec] No se pudo registrar ruta en run_reports")


def generar_desde_run_id(
    run_id: str,
    ruta_db: str = "motor_bt.db"
) -> Tuple[str, str]:
    try:
        runs = obtener_ejecuciones(ruta_db=ruta_db)
        datos_run = next((r for r in runs if r.get("run_id") == run_id), None)
        if not datos_run:
            return ("", "")

        circuitos = obtener_circuitos(run_id=run_id, ruta_db=ruta_db)
        if not circuitos:
            return ("", "")

        ruta_base = os.getcwd()
        ruta_docx = generar_memoria_docx(datos_run, circuitos, ruta_base)
        ruta_pdf = generar_reporte_pdf(datos_run, circuitos, ruta_base)

        _registrar_ruta_reporte(run_id, "DOCX", ruta_docx, ruta_db)
        _registrar_ruta_reporte(run_id, "PDF", ruta_pdf, ruta_db)
        return (ruta_docx, ruta_pdf)
    except Exception as e:
        print(f"[reporteria_sec] Error regenerando reportes: {e}")
        return ("", "")
